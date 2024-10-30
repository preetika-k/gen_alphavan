
"""import requests
from docx import Document
from docx.shared import Pt

# Define your API key and other parameters
api_key = "ME8DM10PKN7RSIM6"
symbol = "IBM"
function = "TIME_SERIES_DAILY"  # Alpha Vantage function for daily stock data

# Construct the API URL
url = f"https://www.alphavantage.co/query?function={function}&symbol={symbol}&apikey={api_key}"

# Make the API request
response = requests.get(url)

# Check if the request was successful
if response.status_code == 200:
    data = response.json()
    #print(data)
    time_series = data.get("Time Series (Daily)", {})
    highs = []
    lows = []
    for date, metrics in sorted(time_series.items(), reverse=True)[:10]:
        high = float(metrics["2. high"])
        low = float(metrics["3. low"])
        highs.append(high)
        lows.append(low)

# Print the results
    print("Last 10 days' highs:", highs)
    print("Last 10 days' lows:", lows)
    

else:
    print("Error:", response.status_code)
""" 
import requests
from docx import Document
from docx.shared import Pt

# Define your API key and other parameters
api_key = "ME8DM10PKN7RSIM6"
symbol = "IBM"
function = "TIME_SERIES_DAILY"  # Alpha Vantage function for daily stock data

# Construct the API URL
url = f"https://www.alphavantage.co/query?function={function}&symbol={symbol}&apikey={api_key}"

# Make the API request
response = requests.get(url)

# Check if the request was successful
if response.status_code == 200:
    data = response.json() 
    time_series = data.get("Time Series (Daily)", {})
    
    # Initialize lists to store dates, highs, and lows for the last 10 days
    last_10_days_data = sorted(time_series.items(), reverse=True)[:10]

    # Create a new Word document
    doc = Document()
    doc.add_heading(f'{symbol} - Last 10 Days High and Low Prices', level=1)

    # Add a table with headers
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Date"
    hdr_cells[1].text = "High"
    hdr_cells[2].text = "Low"
    
    # Fill the table with the last 10 days of data
    for date, metrics in last_10_days_data:
        high = metrics["2. high"]
        low = metrics["3. low"]
        row_cells = table.add_row().cells
        row_cells[0].text = date
        row_cells[1].text = high
        row_cells[2].text = low

    # Customize the header row (optional)
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)

    # Save the document
    doc.save("Last_10_Days_Highs_and_Lows.docx")
    print("Document created successfully!")

else:
    print("Error:", response.status_code)
